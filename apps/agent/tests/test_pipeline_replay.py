"""End-to-end run of the analysis pipeline, offline.

Every LLM call is served from the committed fixtures and every document is read
from the committed dataroom, so this exercises the real use case — preparation,
extraction, merge, conflict resolution, synthesis, scoring — with no API key, no
bucket and no network. It is also the test that catches a prompt edit: change a
prompt and the fixture key changes, replay misses, and this fails loudly instead
of the demo quietly degrading.
"""

import pathlib

import pytest

from src.core.config import settings
from src.data.analyze import document_renderer
from src.domain.analyze.entities import AnalyzeInput, Document
from src.domain.analyze.use_cases import AnalyzeUseCase

ROOT = pathlib.Path(__file__).resolve().parents[1]
DATAROOM = ROOT / "fixtures" / "dataroom"
COMPANY = "Northwind Robotics"
COMPANY_ID = "00000000-0000-4000-8000-000000000002"
AUTOMATION_ID = "00000000-0000-4000-8000-000000000001"


@pytest.fixture
def pipeline(tmp_path, monkeypatch):
    monkeypatch.setattr(settings, "llm_driver", "replay")
    monkeypatch.setattr(settings, "llm_fixture_dir", str(ROOT / "fixtures" / "llm"))
    monkeypatch.setattr(settings, "storage_driver", "local")
    monkeypatch.setattr(settings, "storage_local_root", str(tmp_path))
    monkeypatch.setattr(settings, "google_cloud_bucket_name", "local-bucket")

    # LibreOffice renders the PDF and is not present everywhere; the DOCX render
    # before it still runs, so template breakage is still caught.
    async def _no_pdf(docx_bytes):
        return b"%PDF-1.4 stub"

    monkeypatch.setattr(document_renderer, "convert_docx_to_pdf", _no_pdf)
    monkeypatch.setattr(
        "src.domain.analyze.use_cases.convert_docx_to_pdf", _no_pdf
    )

    from src.data.storage import LocalStorage

    storage = LocalStorage()
    documents = []
    for i, path in enumerate(sorted(DATAROOM.iterdir()), start=1):
        key = f"{COMPANY_ID}/{AUTOMATION_ID}/{path.name}"
        storage.upload_bytes(key, path.read_bytes(), "application/octet-stream")
        documents.append(
            Document(
                id=f"00000000-0000-4000-8000-00000000010{i}",
                url=f"gs://local-bucket/{key}",
            )
        )
    return documents


async def run(documents):
    return await AnalyzeUseCase().execute(
        AnalyzeInput(
            company_id=COMPANY_ID,
            company_name=COMPANY,
            automation_id=AUTOMATION_ID,
            documents=documents,
        )
    )


async def test_the_whole_pipeline_runs_with_no_network(pipeline):
    pdf_url, _, merged, _ = await run(pipeline)

    assert pdf_url == f"gs://local-bucket/one-pagers/{AUTOMATION_ID}.pdf"
    assert merged.facts, "extraction produced nothing"


async def test_the_planted_revenue_disagreement_is_caught(pipeline):
    """Three documents state FY2024 revenue three ways. Finding that is the
    product; a pipeline that reports one number and moves on has failed."""
    _, _, merged, _ = await run(pipeline)

    conflict = next(c for c in merged.conflicts if c.field == "annual_revenue_fy2024")
    joined = " ".join(conflict.values)

    assert "£4.1M" in joined and "£3.8M" in joined and "£3.2M" in joined

    # Computed, not canned. This assertion used to pass because a hand-written
    # dict in record_demo_fixtures.py returned the string "£3.2M" and the
    # fixture round-tripped it — it proved a hash worked, not that anything
    # adjudicated. The rules now run: two of the three values are pro-forma and
    # only the audited accounts state an actual, so no model call is involved in
    # settling it at all.
    assert conflict.preferred_value == "£3.2M"
    assert conflict.preferred_source == "04_audited_accounts.pdf"
    assert conflict.resolution_basis == "source_type"
    assert conflict.confidence == 1.0
    assert conflict.magnitude == "28% spread, £3.2M to £4.1M"


async def test_the_memorandum_prints_the_figure_the_rule_chose(pipeline):
    """The last hop, and the one nothing covered.

    Adjudication reaches synthesis as advisory prompt text and a model writes
    the headline. Every test above can pass — the disagreement found, the
    audited actual chosen, the reason recorded, the conflict view correct —
    while the memorandum itself prints the pitch deck's number.
    """
    _, _, merged, one_pager = await run(pipeline)

    conflict = next(c for c in merged.conflicts if c.field == "annual_revenue_fy2024")
    revenue = one_pager.financial_highlights.annual_revenue

    assert conflict.preferred_value in revenue, (
        f"the rule chose {conflict.preferred_value}; the memorandum reads {revenue!r}"
    )
    assert "£4.1M" not in revenue and "£3.8M" not in revenue


async def test_the_memorandum_reports_the_disagreement_it_resolved(pipeline, monkeypatch):
    """The test above proves the memo prints the right number. This one proves
    it says the other two existed.

    A memorandum that states £3.2M and stops is indistinguishable from one
    produced by a pipeline that never saw the deck or the model, which is the
    difference this product is sold on. The document that reaches an investment
    committee has to carry the disagreement, not just survive it.
    """
    import io

    import docx

    captured = {}

    async def _capture(docx_bytes):
        captured["docx"] = docx_bytes
        return b"%PDF-1.4 stub"

    monkeypatch.setattr("src.domain.analyze.use_cases.convert_docx_to_pdf", _capture)

    _, _, merged, _ = await run(pipeline)

    document = docx.Document(io.BytesIO(captured["docx"]))
    text = "\n".join(p.text for p in document.paragraphs)

    conflict = next(c for c in merged.conflicts if c.field == "annual_revenue_fy2024")
    assert conflict.preferred_value in text
    for rejected in ("£4.1M", "£3.8M"):
        assert rejected in text, (
            f"the dataroom states {rejected} and the memorandum never mentions it"
        )
    assert conflict.resolution_basis in text, "the memo states a figure without the rule behind it"


async def test_facts_keep_the_document_they_came_from(pipeline):
    _, _, merged, _ = await run(pipeline)

    revenue = merged.facts["annual_revenue_fy2024"]
    sources = {f.source for f in revenue}

    assert any("pitch_deck" in s for s in sources)
    assert any("audited_accounts" in s for s in sources)
    assert all(f.quote for f in revenue), "a fact with no quote cannot be checked"


async def test_every_quote_is_checked_against_the_document_it_cites(pipeline):
    """A citation nobody checked is decoration.

    The PDF half of this is the part that used to be impossible: Step 0 uploads
    each PDF and keeps only the file_id, so by extraction time there were no
    bytes left to check a quote against. The text layer now rides along.
    """
    _, _, merged, _ = await run(pipeline)
    facts = [f for fs in merged.facts.values() for f in fs]

    assert all(f.grounding for f in facts), "every fact is classified"
    assert all(f.quote_verified is not None for f in facts), (
        "every document in this dataroom is machine-readable, so no fact should "
        "come back unverifiable"
    )

    for name in ("01_pitch_deck.pdf", "04_audited_accounts.pdf"):
        from_pdf = [f for f in facts if f.source == name]
        assert from_pdf and all(f.quote_verified for f in from_pdf), (
            f"{name}: quotes should be found verbatim in the PDF text layer"
        )


async def test_no_fact_cites_a_sheet_it_did_not_come_from(pipeline):
    """Excel is extracted one prepared document per sheet, and the canned answers
    used to be selected by file name — so both sheets of a workbook were given the
    same facts, and the Summary sheet's revenue was also recorded against the
    Pipeline sheet, which holds nothing but sales accounts.

    Beyond the wrong page reference, the same fact then arrived from two
    apparently different sources, and anything counting sources reads one sheet
    double-counted as two documents agreeing.
    """
    _, _, merged, _ = await run(pipeline)
    facts = [f for fs in merged.facts.values() for f in fs]

    misattributed = [f for f in facts if f.quote_verified is False]
    assert not misattributed, [
        f"{f.source}: {f.field} = {f.quote!r}" for f in misattributed
    ]


async def test_the_scorecard_survives_the_pipeline(pipeline):
    """The one-pager is now a return value, not just a rendered PDF.

    Everything the frontend will show — the weighted scorecard, the coverage
    denominator, the summary — used to exist for the length of one function call
    and then be discarded in favour of a URL.
    """
    _, _, _, one_pager = await run(pipeline)

    assert len(one_pager.scorecard) == 8
    assert all(c.weighted_score for c in one_pager.scorecard)
    assert one_pager.overall_score
    assert one_pager.executive_summary


async def test_a_difference_of_basis_is_not_reported_as_a_contradiction(pipeline):
    """Year-end headcount of 52 and a 49 average are both true. Flagging that as
    a conflict would bury the revenue one in noise."""
    _, _, merged, _ = await run(pipeline)

    assert not any(c.field == "employees" for c in merged.conflicts)


async def test_absent_information_is_reported_rather_than_assumed(pipeline):
    """The dataroom has no contracts, insurance or policies. Saying so is what
    turns the run into a document request list."""
    _, _, merged, _ = await run(pipeline)

    assert "cap_table" in merged.coverage
    assert {"insurance", "policies", "client_contracts"} <= set(merged.missing)


async def test_a_changed_prompt_invalidates_the_fixtures(pipeline, monkeypatch):
    """Guards the guard: if replay silently fell back to something, none of the
    tests above would mean anything."""
    from src.core import llm
    from src.core.prompts import one_pager as one_pager_prompts

    monkeypatch.setattr(
        one_pager_prompts, "ONE_PAGER_SYSTEM_PROMPT", "changed {current_date}"
    )
    monkeypatch.setattr(
        "src.data.analyze.one_pager_service.ONE_PAGER_SYSTEM_PROMPT",
        "changed {current_date}",
    )

    with pytest.raises(llm.ReplayMiss):
        await run(pipeline)


# ---------------------------------------------------------------------------
# The four domain reports. Until these fixtures existed, diligence_extraction.py
# and diligence_synthesis.py — some 600 lines behind the four reports advertised
# on the README front page — had no offline coverage at all, and stage 2 of
# `make demo` died on a replay miss.
# ---------------------------------------------------------------------------

DOMAINS = ["OPERATIONAL", "COMMERCIAL", "FINANCIAL", "CAP_TABLE_AND_LEGAL_REVIEW"]


@pytest.mark.parametrize("index, domain", list(enumerate(DOMAINS)))
async def test_each_domain_report_runs_with_no_network(pipeline, monkeypatch, index, domain):
    from src.data.diligence import document_renderer as diligence_renderer
    from src.domain.diligence.entities import DiligenceInput
    from src.domain.diligence.use_cases import DiligenceUseCase

    async def _no_pdf(docx_bytes):
        return b"%PDF-1.4 stub"

    monkeypatch.setattr(diligence_renderer, "convert_docx_to_pdf", _no_pdf, raising=False)
    monkeypatch.setattr("src.domain.diligence.use_cases.convert_docx_to_pdf", _no_pdf)

    # The same child automation ids the recorder used, and the ones the backend
    # creates for stage 2. A shared id would have each domain overwrite the
    # previous one's cached facts.
    automation_id = f"{AUTOMATION_ID[:-1]}{index + 2}"

    url = await DiligenceUseCase(domain).execute(
        DiligenceInput(
            company_id=COMPANY_ID,
            company_name=COMPANY,
            automation_id=automation_id,
            domain=domain,
            documents=pipeline,
        )
    )

    assert url.startswith("gs://local-bucket/")
    assert url.endswith(".pdf")


async def test_the_financial_report_carries_the_disagreement_into_stage_two(pipeline, monkeypatch):
    """The conflict found in triage has to survive into the domain report, or
    stage 2 quietly restates whichever figure it happened to read last."""
    from src.data.diligence import document_renderer as diligence_renderer
    from src.data.diligence.report_service import DiligenceReportService
    from src.domain.diligence.entities import DiligenceInput
    from src.domain.diligence.use_cases import DiligenceUseCase

    async def _no_pdf(docx_bytes):
        return b"%PDF-1.4 stub"

    monkeypatch.setattr(diligence_renderer, "convert_docx_to_pdf", _no_pdf, raising=False)
    monkeypatch.setattr("src.domain.diligence.use_cases.convert_docx_to_pdf", _no_pdf)

    captured = {}
    original = DiligenceReportService.generate

    async def spy(self, domain, company_name, merged):
        captured["merged"] = merged
        return await original(self, domain, company_name, merged)

    monkeypatch.setattr(DiligenceReportService, "generate", spy)

    await DiligenceUseCase("FINANCIAL").execute(
        DiligenceInput(
            company_id=COMPANY_ID, company_name=COMPANY,
            automation_id=f"{AUTOMATION_ID[:-1]}4",
            domain="FINANCIAL", documents=pipeline,
        )
    )

    merged = captured["merged"]
    revenue = merged.facts["annual_revenue_fy2024"]
    assert {f.value for f in revenue} == {"£3.2M", "£3.8M"}
    assert any(c.field == "annual_revenue_fy2024" for c in merged.conflicts)
